"""
Word document parser using python-docx.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from ..models import Document, compute_file_hash
from ..exceptions import ParseError
from .base import BaseParser

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Word document parser using python-docx."""
    
    def supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        return [".docx", ".doc"]
    
    def supported_mimetypes(self) -> List[str]:
        """Return list of supported MIME types."""
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]
    
    def parse(self, file_path: Path) -> Document:
        """
        Parse Word file and return Document.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Parsed Document
            
        Raises:
            ParseError: If parsing fails
        """
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        file_path = file_path.resolve()
        
        if not file_path.exists():
            raise ParseError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ParseError(f"Not a file: {file_path}")
        
        try:
            # Extract content based on file type
            if file_path.suffix.lower() == ".docx":
                content = self._extract_docx(file_path)
            else:
                # For .doc files, try to use textract or antiword as fallback
                content = self._extract_doc_fallback(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(file_path)
            
            # Compute content hash
            content_hash = compute_file_hash(file_path)
            
            return Document(
                file_path=file_path,
                content=content,
                metadata=metadata,
                content_hash=content_hash,
                parsed_at=datetime.now(),
                file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if file_path.suffix.lower() == ".docx"
                else "application/msword",
            )
            
        except Exception as e:
            raise ParseError(f"Failed to parse Word document {file_path}: {e}") from e
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from .docx file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ParseError("python-docx is required for Word document parsing. Install with: pip install python-docx")
        
        doc = DocxDocument(str(file_path))
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if it's a heading
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                    try:
                        heading_level = int(level)
                    except ValueError:
                        heading_level = 1
                    text_parts.append(f"{'#' * heading_level} {text}")
                else:
                    text_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(table_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_table(self, table) -> str:
        """Extract text from a table as markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        
        if len(rows) >= 1:
            # Add header separator after first row
            col_count = len(table.rows[0].cells) if table.rows else 0
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            rows.insert(1, separator)
        
        return "\n".join(rows) if rows else ""
    
    def _extract_doc_fallback(self, file_path: Path) -> str:
        """Fallback extraction for .doc files."""
        # Try antiword first (common on Linux/Mac)
        import subprocess
        
        try:
            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # Try catdoc as another fallback
        try:
            result = subprocess.run(
                ["catdoc", str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # If all else fails, raise an error with helpful message
        raise ParseError(
            f"Cannot parse .doc file {file_path}. "
            "Install 'antiword' or 'catdoc' for .doc support, "
            "or convert to .docx format."
        )
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Word file."""
        metadata = {
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
        }
        
        if file_path.suffix.lower() == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                core_props = doc.core_properties
                
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.keywords:
                    metadata["keywords"] = core_props.keywords
                if core_props.created:
                    metadata["creation_date"] = core_props.created.isoformat()
                if core_props.modified:
                    metadata["modified_date"] = core_props.modified.isoformat()
                
                # Count paragraphs as a rough page estimate
                metadata["paragraph_count"] = len(doc.paragraphs)
                
            except Exception:
                pass
        
        return metadata
